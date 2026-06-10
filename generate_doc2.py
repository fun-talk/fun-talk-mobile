# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SOFTWARE_NAME = "聊天气氛组移动客户端软件 V1.0"
OUTPUT_DIR = "软著材料"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "源代码文档_聊天气氛组移动客户端软件_V1.0.docx")
LINES_PER_PAGE = 55
FRONT_PAGES = 30
BACK_PAGES = 30

files = [
    "src/app/_layout.tsx",
    "src/app/index.tsx",
    "src/app/(auth)/_layout.tsx",
    "src/app/(auth)/login.tsx",
    "src/app/(app)/_layout.tsx",
    "src/app/(app)/courses/index.tsx",
    "src/app/(app)/lesson.tsx",
    "src/features/auth/AuthProvider.tsx",
    "src/features/auth/components/LoginScreen.tsx",
    "src/features/auth/components/LoginView.tsx",
    "src/features/auth/services/login.ts",
    "src/features/auth/hooks/useWechatQrLogin.ts",
    "src/features/auth/components/WechatModal.tsx",
    "src/features/courses/components/CourseHomeScreen.tsx",
    "src/features/courses/components/CourseMapBackground.tsx",
    "src/features/courses/components/CourseNode.tsx",
    "src/features/courses/hooks/useCourseHome.ts",
    "src/features/courses/services/courseHomeApi.ts",
    "src/features/courses/layout/courseHomeLayout.ts",
    "src/features/lesson/LessonScreen.tsx",
    "src/features/lesson/components/LessonWebViewScreen.tsx",
    "src/features/lesson/components/NativeLessonScreen.tsx",
    "src/features/lesson/components/NativeLessonShell.tsx",
    "src/features/lesson/components/FreeChatPanel.tsx",
    "src/features/lesson/components/RecordingPanel.tsx",
    "src/features/lesson/nativeLessonController.ts",
    "src/features/lesson/nativeLessonMedia.ts",
    "src/features/lesson/nativeLessonRealtimeProjection.ts",
    "src/features/lesson/nativeLessonLoader.ts",
    "src/features/lesson/nativeLessonProgress.ts",
    "src/features/lesson/recordingController.ts",
    "src/features/lesson/simpleVad.ts",
    "src/features/lesson/webViewBootstrap.ts",
    "src/features/lesson/webViewMessages.ts",
    "src/features/lesson/buildLessonWebUrl.ts",
    "src/features/lesson/hooks/useNativeLessonController.ts",
    "src/features/lesson/hooks/useNativeLessonRecording.ts",
    "src/features/lesson/hooks/useNativeLessonRealtimeSession.ts",
    "src/lib/api/client.ts",
    "src/lib/auth/session.ts",
    "src/lib/auth/storage.ts",
    "src/lib/auth/storageCore.ts",
    "src/lib/device/deviceId.ts",
    "src/lib/env.ts",
    "src/shared/courseHomeMap.ts",
    "src/shared/courseHomeProgress.ts",
    "src/shared/courseOpeningSceneEntry.ts",
    "src/hooks/use-theme.ts",
    "src/components/OpeningAnimation.tsx",
    "src/components/themed-text.tsx",
    "src/components/themed-view.tsx",
]

def is_block_end(line):
    s = line.rstrip()
    if not s:
        return False
    if s == '}' or s.startswith('}') or s.endswith('}'):
        return True
    if s in (');', ');', ']'):
        return True
    return False

def find_cut(lines, target, direction='forward'):
    if direction == 'forward':
        for i in range(target, min(len(lines), target + 300)):
            if lines[i].startswith('// ===== File:'):
                return i
            if is_block_end(lines[i]) and i > target:
                return i + 1
        return min(len(lines), target + 150)
    else:
        for i in range(target, max(0, target - 300), -1):
            if lines[i].startswith('// ===== File:'):
                return i
            if is_block_end(lines[i]) and i < target:
                return i + 1
        return max(0, target - 150)

def add_page_number(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def collect_code():
    all_lines = []
    stats = []
    for f in files:
        if not os.path.exists(f):
            print("MISSING:", f)
            continue
        with open(f, 'r', encoding='utf-8') as fh:
            content = fh.read()
        raw = content.split('\n')
        cleaned = []
        prev_blank = False
        for line in raw:
            is_blank = line.strip() == ''
            if is_blank and prev_blank:
                continue
            cleaned.append(line)
            prev_blank = is_blank
        stats.append((f, len(cleaned)))
        all_lines.append("// ===== File: " + f + " =====")
        all_lines.extend(cleaned)
        all_lines.append('')
        all_lines.append('')
    return all_lines, stats

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    all_lines, stats = collect_code()
    total = len(all_lines)
    estimated_pages = total // LINES_PER_PAGE + (1 if total % LINES_PER_PAGE else 0)
    print("Total cleaned lines:", total)
    print("Estimated pages ({} lines/page):".format(LINES_PER_PAGE), estimated_pages)

    front_target = FRONT_PAGES * LINES_PER_PAGE
    back_target = BACK_PAGES * LINES_PER_PAGE

    if total <= front_target + back_target:
        trimmed = all_lines
        was_trimmed = False
        front_cut = 0
        back_kept = total
    else:
        front_cut = find_cut(all_lines, front_target, 'forward')
        back_start = find_cut(all_lines, total - back_target, 'backward')
        if front_cut >= back_start:
            trimmed = all_lines
            was_trimmed = False
            front_cut = 0
            back_kept = total
        else:
            trimmed = all_lines[:front_cut] + all_lines[back_start:]
            was_trimmed = True
            back_kept = total - back_start

    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    header = section.header
    header.is_linked_to_previous = False
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.add_run(SOFTWARE_NAME)
    header_run.font.name = '宋体'
    header_run.font.size = Pt(10.5)
    header_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_page_number(section)

    buffer_lines = []
    for line in trimmed:
        if line.startswith('// ===== File:'):
            if buffer_lines:
                para = doc.add_paragraph('\n'.join(buffer_lines))
                pf = para.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
                for run in para.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10.5)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
                buffer_lines = []
            para = doc.add_paragraph(line)
            pf = para.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            for run in para.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(10.5)
                run.font.bold = True
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
        else:
            buffer_lines.append(line)

    if buffer_lines:
        para = doc.add_paragraph('\n'.join(buffer_lines))
        pf = para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        for run in para.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(10.5)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')

    doc.save(OUTPUT_FILE)
    print("\nDocument saved to:", OUTPUT_FILE)
    print("Lines in document:", len(trimmed))
    print("Was trimmed:", was_trimmed)
    if was_trimmed:
        print("Front part: first", front_cut, "lines")
        print("Back part: last", back_kept, "lines")
        print("Middle discarded:", total - front_cut - back_kept, "lines")

    print("\n===== REPORT =====")
    print("Files processed:", len(stats))
    print("Original total lines (cleaned):", sum(c for _, c in stats))
    print("Lines in final document:", len(trimmed))
    final_est = len(trimmed) // LINES_PER_PAGE + (1 if len(trimmed) % LINES_PER_PAGE else 0)
    print("Estimated total pages ({} lines/page):".format(LINES_PER_PAGE), final_est)
    print("Trimmed (60-page limit):", "Yes" if was_trimmed else "No")

if __name__ == '__main__':
    main()
