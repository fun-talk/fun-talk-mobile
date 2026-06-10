import os
import re
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ===== Configuration =====
SOFTWARE_NAME = "聊天气氛组移动客户端软件 V1.0"
OUTPUT_DIR = "软著材料"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "源代码文档_聊天气氛组移动客户端软件_V1.0.docx")
LINES_PER_PAGE_TARGET = 50  # minimum lines per page
TARGET_PAGES = 60
FRONT_PAGES = 30
BACK_PAGES = 30

# File list in logical order
files = [
    "src/app/_layout.tsx",
    "src/app/index.tsx",
    "src/app/(auth)/_layout.tsx",
    "src/app/(auth)/login.tsx",
    "src/app/(app)/_layout.tsx",
    "src/app/(app)/courses/index.tsx",
    "src/app/(app)/lesson.tsx",
    "src/features/auth/AuthProvider.tsx",
    "src/features/auth/components/LoginScreen.tsx",
    "src/features/auth/components/LoginView.tsx",
    "src/features/auth/services/login.ts",
    "src/features/auth/hooks/useWechatQrLogin.ts",
    "src/features/auth/components/WechatModal.tsx",
    "src/features/courses/components/CourseHomeScreen.tsx",
    "src/features/courses/components/CourseMapBackground.tsx",
    "src/features/courses/components/CourseNode.tsx",
    "src/features/courses/hooks/useCourseHome.ts",
    "src/features/courses/services/courseHomeApi.ts",
    "src/features/courses/layout/courseHomeLayout.ts",
    "src/features/lesson/LessonScreen.tsx",
    "src/features/lesson/components/LessonWebViewScreen.tsx",
    "src/features/lesson/components/NativeLessonScreen.tsx",
    "src/features/lesson/components/NativeLessonShell.tsx",
    "src/features/lesson/components/FreeChatPanel.tsx",
    "src/features/lesson/components/RecordingPanel.tsx",
    "src/features/lesson/nativeLessonController.ts",
    "src/features/lesson/nativeLessonMedia.ts",
    "src/features/lesson/nativeLessonRealtimeProjection.ts",
    "src/features/lesson/nativeLessonLoader.ts",
    "src/features/lesson/nativeLessonProgress.ts",
    "src/features/lesson/recordingController.ts",
    "src/features/lesson/simpleVad.ts",
    "src/features/lesson/webViewBootstrap.ts",
    "src/features/lesson/webViewMessages.ts",
    "src/features/lesson/buildLessonWebUrl.ts",
    "src/features/lesson/hooks/useNativeLessonController.ts",
    "src/features/lesson/hooks/useNativeLessonRecording.ts",
    "src/features/lesson/hooks/useNativeLessonRealtimeSession.ts",
    "src/lib/api/client.ts",
    "src/lib/auth/session.ts",
    "src/lib/auth/storage.ts",
    "src/lib/auth/storageCore.ts",
    "src/lib/device/deviceId.ts",
    "src/lib/env.ts",
    "src/shared/courseHomeMap.ts",
    "src/shared/courseHomeProgress.ts",
    "src/shared/courseOpeningSceneEntry.ts",
    "src/hooks/use-theme.ts",
    "src/components/OpeningAnimation.tsx",
    "src/components/themed-text.tsx",
    "src/components/themed-view.tsx",
]

# ===== Helper functions =====
def is_code_block_end(line):
    """Heuristic: a line ending with } at root level or file end marks a block end."""
    stripped = line.rstrip()
    if not stripped:
        return False
    if stripped == '}' or stripped.startswith('}') or stripped.endswith('}'):
        return True
    if stripped in (');', ');', ']'):
        return True
    return False

def find_good_cut(lines, target_index, direction='forward'):
    """Find a good cut point near target_index."""
    if direction == 'forward':
        for i in range(target_index, min(len(lines), target_index + 200)):
            if lines[i].startswith('// ===== File:'):
                return i
            if is_code_block_end(lines[i]) and i > target_index:
                return i + 1
        return min(len(lines), target_index + 100)
    else:
        for i in range(target_index, max(0, target_index - 200), -1):
            if lines[i].startswith('// ===== File:'):
                return i
            if is_code_block_end(lines[i]) and i < target_index:
                return i + 1
        return max(0, target_index - 100)

def add_page_number(section):
    """Add page number to the footer, aligned right."""
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def set_paragraph_format(paragraph, font_name='Courier New', font_size=Pt(10.5)):
    """Set paragraph to monospace, small margins, single spacing."""
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.line_spacing = 1.0
    
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = font_size
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def collect_code():
    all_lines = []
    file_stats = []
    for f in files:
        if not os.path.exists(f):
            print(f"MISSING: {f}")
            continue
        with open(f, 'r', encoding='utf-8') as fh:
            content = fh.read()
        raw_lines = content.split('\n')
        cleaned = []
        prev_blank = False
        for line in raw_lines:
            is_blank = line.strip() == ''
            if is_blank and prev_blank:
                continue
            cleaned.append(line)
            prev_blank = is_blank
        
        file_stats.append((f, len(cleaned)))
        all_lines.append(f"// ===== File: {f} =====")
        all_lines.extend(cleaned)
        all_lines.append('')
        all_lines.append('')
    return all_lines, file_stats

def trim_to_pages(lines, front_lines_target, back_lines_target):
    """Trim lines to front + back pages, cutting at file boundaries or block ends."""
    if len(lines) <= front_lines_target + back_lines_target:
        return lines, False, 0, 0
    
    front_cut = find_good_cut(lines, front_lines_target, 'forward')
    back_start = find_good_cut(lines, len(lines) - back_lines_target, 'backward')
    
    if front_cut >= back_start:
        return lines, False, 0, 0
    
    trimmed = lines[:front_cut] + lines[back_start:]
    return trimmed, True, front_cut, len(lines) - back_start

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    all_lines, file_stats = collect_code()
    total_lines = len(all_lines)
    print(f"Total cleaned lines: {total_lines}")
    
    estimated_pages = total_lines // LINES_PER_PAGE_TARGET + (1 if total_lines % LINES_PER_PAGE_TARGET else 0)
    print(f"Estimated pages (min 50 lines/page): {estimated_pages}")
    
    front_target = FRONT_PAGES * LINES_PER_PAGE_TARGET
    back_target = BACK_PAGES * LINES_PER_PAGE_TARGET
    
    trimmed_lines, was_trimmed, front_cut, back_kept = trim_to_pages(all_lines, front_target, back_target)
    
    doc = Document()
    
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    
    # Header
    header = section.header
    header.is_linked_to_previous = False
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.add_run(SOFTWARE_NAME)
    header_run.font.name = '宋体'
    header_run.font.size = Pt(10.5)
    header_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # Footer with page number
    add_page_number(section)
    
    # Add code lines
    buffer_lines = []
    for line in trimmed_lines:
        if line.startswith('// ===== File:'):
            if buffer_lines:
                para = doc.add_paragraph('\n'.join(buffer_lines))
                set_paragraph_format(para)
                buffer_lines = []
            para = doc.add_paragraph(line)
            set_paragraph_format(para)
            para.runs[0].font.bold = True
        else:
            buffer_lines.append(line)
    
    if buffer_lines:
        para = doc.add_paragraph('\n'.join(buffer_lines))
        set_paragraph_format(para)
    
    doc.save(OUTPUT_FILE)
    print(f"\nDocument saved to: {OUTPUT_FILE}")
    print(f"Lines in document: {len(trimmed_lines)}")
    print(f"Was trimmed: {was_trimmed}")
    if was_trimmed:
        print(f"Front part: first {front_cut} lines")
        print(f"Back part: last {back_kept} lines")
        print(f"Middle discarded: {total_lines - front_cut - back_kept} lines")
    
    print("\n===== REPORT =====")
    print(f"Files processed: {len(file_stats)}")
    total_raw = sum(c for _, c in file_stats)
    print(f"Original total lines (cleaned): {total_raw}")
    print(f"Lines in final document: {len(trimmed_lines)}")
    print(f"Estimated total pages (min 50/page): {len(trimmed_lines) // LINES_PER_PAGE_TARGET + (1 if len(trimmed_lines) % LINES_PER_PAGE_TARGET else 0)}")
    print(f"Trimmed (60-page limit): {'Yes' if was_trimmed else 'No'}")

if __name__ == '__main__':
    main()
